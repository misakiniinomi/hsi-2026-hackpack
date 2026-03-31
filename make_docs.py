from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p

def add_bullet(doc, items, symbol='•'):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.text = item
        p.runs[0].font.size = Pt(10)

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.text = item
        p.runs[0].font.size = Pt(10)

def add_divider(doc):
    doc.add_paragraph('─' * 60)

def add_two_col_section(doc, left_title, left_items, right_title, right_items):
    # We can't do true 2-col in simple docx, so present as two labeled blocks
    p = doc.add_paragraph()
    p.add_run(f'◆ {left_title}').bold = True
    for item in left_items:
        doc.add_paragraph(f'  ✗  {item}', style='List Bullet')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run(f'◆ {right_title}').bold = True
    for item in right_items:
        doc.add_paragraph(f'  ✓  {item}', style='List Bullet')

# ═══════════════════════════════════════════════════════════
# BGC
# ═══════════════════════════════════════════════════════════
def make_bgc():
    doc = Document()

    # Title
    title = doc.add_heading('Boys & Girls Club — Youth Engagement', 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 48, 135)

    doc.add_heading('Help more Rainier Valley families discover and join the Club.', 1)
    add_body(doc, 'Hack for Social Impact Seattle 2026  |  April 26, 2026  |  AI House, Seattle')
    add_body(doc, 'Rainier Vista Boys & Girls Club has life-changing programs. The problem isn\'t the programs — it\'s that families don\'t know the door is open.')

    add_divider(doc)

    # About the Club
    doc.add_heading('About the Club', 1)
    doc.add_heading('Smilow Rainier Vista Boys & Girls Club', 2)
    add_body(doc, 'A nonprofit youth development facility serving K–12 kids and teens in Rainier Valley, Seattle. After-school programs, sports, homework help, cooking, robotics, career training — and a community that becomes a second home.')
    doc.add_paragraph()
    add_body(doc, 'Key Info:')
    add_bullet(doc, [
        '📍 4520 Martin Luther King Jr Way S, Seattle, WA 98108',
        '🌐 positiveplace.org/clubs/smilow-rainier-vista',
        '🏟️  40,000 sq ft — gym, music studio, tech lab, café',
        '💵 Annual membership: $50 (scholarships available)',
        '🕑 Mon–Fri, 2–6pm (school year)  ·  Summer: up to 12h/day, 70 days',
    ])
    doc.add_paragraph()
    add_body(doc, 'What\'s actually happening inside:')
    add_bullet(doc, [
        '186 teens earned CPR / First Aid in 2025',
        '146 teens got paid work experience through Club partnerships',
        '2,294 free meals served last summer',
        '24 members enrolled in 4-year universities after graduating',
    ])

    add_divider(doc)

    # The Challenge
    doc.add_heading('The Challenge', 1)
    doc.add_heading('How do we get families to know about it? How do we make them want to come?', 2)
    add_body(doc, 'This is where the real battle is.')
    doc.add_paragraph()
    add_two_col_section(
        doc,
        'Today: Great programs, invisible club',
        [
            'After-school & teen enrollment declining for 2 years',
            'Very limited marketing budget, no dedicated outreach staff',
            'Currently uses Canva for flyers — but staff aren\'t comfortable with it either, and output is infrequent',
            '225 kids participated last year but never joined',
        ],
        'Your goal: Close that gap',
        [
            'Families in Rainier Valley learn the Club exists',
            'They feel safe and excited to try it',
            'The first step — visit or sign-up — feels easy',
            'Staff can market without a marketing team',
        ]
    )

    add_divider(doc)

    # The Landscape
    doc.add_heading('The Landscape', 1)
    doc.add_heading('4 competitors within walking distance — 3 of them free', 2)
    add_body(doc, 'B&G Club charges $50/year. Every nearby alternative is free or near-free. YMCA has direct school partnerships. This is the uphill battle your solution must navigate.')
    doc.add_paragraph()
    competitors = [
        'Rainier Community Center (0.4 mi) — Seattle Parks, free. After-school, sports, art.',
        'Rainier Beach Community Center — Seattle Parks, free. Late-night programs (Fri–Sat).',
        'Jefferson Community Center (~2.5 mi) — Seattle Parks, free–low cost.',
        'YMCA Rainier Valley (0.05 mi) — Fee-based. Direct partnerships with 22 schools. "100% college-bound" track record. Black Achievers program.',
    ]
    add_numbered(doc, competitors)
    doc.add_paragraph()
    add_body(doc, 'Nearby Schools — Potential Outreach Targets (within 3 mi):')
    add_bullet(doc, [
        'Cleveland High School (1.22 mi)',
        'Aki Kurose Middle School (1.25 mi)',
        'Graham Hill Elementary (1.65 mi)',
        'Wing Luke Elementary (2.23 mi)',
        'South Shore PreK-8 (2.85 mi)',
        'Rainier Beach High School (2.87 mi)',
        '+ Dunlap Elementary, MLK Jr. Elementary, Orca K–8, Alan T. Sugiyama HS, South Lake HS',
    ])

    add_divider(doc)

    # Blockers
    doc.add_heading("What's Blocking Awareness", 1)
    doc.add_heading('Four walls — any one of them keeps a family out', 2)
    add_body(doc, 'The programs exist. The problem is the experience of finding them — and feeling like they\'re for you.')
    doc.add_paragraph()
    blockers = [
        ('No way to discover the Club', 'Flyers reach a few. Word of mouth is slow. There is no consistent, digital way for a Rainier Valley family to stumble across the Club.'),
        ('Hard to picture what happens inside', '"We have robotics" means nothing. Families can\'t visualize a typical Tuesday at 4:30pm. The story isn\'t being told — anywhere.'),
        ('Sign-up feels like a commitment before they\'ve tried', 'The national enrollment page is confusing. Many families get stuck or give up before completing registration — even when they\'re interested.'),
        ('SNS is essentially non-existent', 'There is only 1 marketing person covering 30+ clubs across the organization. Club staff are coaches and educators — outreach and marketing are not their job.'),
    ]
    for i, (title, desc) in enumerate(blockers, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        doc.add_paragraph(desc)

    add_divider(doc)

    # Journey
    doc.add_heading('Current Journey: Now vs. What You\'ll Design', 1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❌ Today').bold = True
    today = [
        'Sees a flyer at a community event — rare, inconsistent',
        'Hears from a friend — word of mouth is slow',
        'Checks website — hard to picture what actually happens inside',
        'Tries to sign up — gets stuck on confusing national enrollment page',
        'Gives up — or shows up once and never hears back',
    ]
    add_numbered(doc, today)
    add_body(doc, '→ Most families drop off. 225 participated last year without ever becoming members.')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run('✅ Future — What You Build').bold = True
    future = [
        'A staff member shares the Club\'s value in seconds — no marketing skills needed, no extra time',
        'A family discovers the Club through something shared — SNS, search, or a trusted person',
        'Explores content that shows exactly what happens inside — and feels like it\'s for their kid',
        'Takes an easy first step — visits, scans a QR, or follows a simple sign-up guide',
    ]
    add_numbered(doc, future)
    add_body(doc, '→ Staff are coaches and educators — not marketers. Your solution must work for them, not ask more of them.')

    add_divider(doc)

    # Success
    doc.add_heading('Success', 1)
    doc.add_heading('If a kid says "I want to go there" — you\'ve succeeded.', 2)
    doc.add_paragraph()
    add_body(doc, 'Success Metrics:')
    add_bullet(doc, [
        'Afterschool membership +5%',
        'Increased elementary & middle school enrollment — especially basketball and flag football',
        'Staff can market without a marketing team',
    ])
    doc.add_paragraph()
    add_body(doc, 'Test your solution against these:')
    add_bullet(doc, [
        'Could a Club leader use this with zero training?',
        'Does it work on a smartphone?',
        'Can it start small and grow with minimal budget?',
        'Does a kid say "I want to go there" after seeing it?',
    ])

    add_divider(doc)

    # FAQ
    doc.add_heading('FAQ', 1)
    faqs = [
        ('Can we change the national sign-up page?', 'No. The national enrollment page is locked. It cannot be modified.'),
        ('Can we edit the local Club page?', 'Yes, slightly. Minor text and image edits are possible on the local page — no full redesign.'),
        ('Is there a budget?', 'Very limited. A small amount may be available, but large ongoing costs are not realistic. Prefer low-cost or free solutions.'),
        ('Who is the target audience?', 'Elementary and middle school kids (K–8), or their parents. Both are valid targets.'),
        ('Who should be able to use the solution?', 'Club leaders and coaches — even without marketing training. The simpler, the better.'),
        ('What kind of proposals are most welcome?', 'Things that can start small and be used immediately in the field. Practical tools that staff will actually pick up and use.'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(f'Q: {q}').bold = True
        doc.add_paragraph(f'A: {a}')

    doc.save('C:/Users/misakin/DEV/02_personal/hsi-2026-hackpack/bgc-v2/hackpack-bgc.docx')
    print('BGC done')


# ═══════════════════════════════════════════════════════════
# Hopelink
# ═══════════════════════════════════════════════════════════
def make_hopelink():
    doc = Document()

    # Title
    title = doc.add_heading('Hopelink — Transportation Access', 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 122, 110)

    doc.add_heading('Deliver transportation access simply — without frustration.', 1)
    add_body(doc, 'Hack for Social Impact Seattle 2026  |  April 26, 2026  |  AI House, Seattle')
    add_body(doc, 'Washington state has dozens of transportation programs for people who need them. The information exists. The problem is access. Your challenge: make applying feel effortless.')

    add_divider(doc)

    # About
    doc.add_heading('About Hopelink', 1)
    add_body(doc, 'Hopelink is a Washington state nonprofit that helps people in need build stability and self-sufficiency. Their programs span four areas — and transportation is the thread that makes all the others accessible.')
    doc.add_paragraph()
    add_body(doc, 'Four program areas:')
    add_bullet(doc, ['🚗 Transportation', '🏠 Housing', '🍎 Food', '💰 Financial Assistance'])
    doc.add_paragraph()
    add_body(doc, 'Key Links:')
    add_bullet(doc, [
        'Organization: hopelink.org',
        'Find a Ride service: findaride.org',
    ])

    add_divider(doc)

    # Challenge
    doc.add_heading('The Challenge', 1)
    doc.add_heading('How might we make it simple for anyone who needs a ride to find one — and actually receive it?', 2)
    doc.add_paragraph()
    add_two_col_section(
        doc,
        'Phase 1 (Exists): "Find a Ride"',
        [
            'Can search available services',
            'Cannot apply — links out only',
            'Eligibility is unclear before applying',
            'Each service requires separate application',
        ],
        'Phase 2 (Your Challenge): One-click application',
        [
            'Know which services you qualify for instantly',
            'Apply to multiple services at once',
            'Enter data once — never repeat it',
            'Know required documents before starting',
        ]
    )

    add_divider(doc)

    # Blockers
    doc.add_heading("What's Blocking Access", 1)
    doc.add_heading('Seven barriers — any one of them causes most people to quit', 2)
    add_body(doc, 'The service exists. The problem is the experience of reaching it.')
    doc.add_paragraph()
    barriers = [
        ('Separate application for every service', 'There is no unified entry point. Each provider requires its own form, its own process, its own contact.'),
        ('Too many documents required', 'Income verification, doctor\'s notes, residency proof — discovered only mid-application. Most people give up here.'),
        ('No way to know eligibility upfront', 'Users invest time in an application, then find out they don\'t qualify. There\'s no preview, no filter, no guidance.'),
        ('Same information entered over and over', 'Name, address, income, condition — typed again for every single service. Exhausting and error-prone.'),
        ('Every provider has different rules', 'Conditions vary: age thresholds, income brackets, disability definitions, area boundaries. No unified guide exists.'),
        ('Conditions are confusing and inconsistent', 'Even side-by-side, eligibility criteria are written in jargon, inconsistently, and without plain-language summaries.'),
        ('No data sharing between providers', 'Each organization operates in a silo. Users carry 100% of the burden of coordination. The system doesn\'t work for them — they work for the system.'),
    ]
    for i, (title, desc) in enumerate(barriers, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        doc.add_paragraph(desc)
    doc.add_paragraph()
    add_body(doc, '→ Result: Most people abandon mid-process. Those who need help most are least likely to receive it.')

    add_divider(doc)

    # Personas
    doc.add_heading('User Personas', 1)
    doc.add_heading('Real users. Different needs. Same broken system.', 2)
    doc.add_paragraph()
    personas = [
        ('👵 Elderly Adult',
         'Not tech-savvy. Needs transportation to medical appointments. May have vision or mobility challenges. Often relies on others to navigate digital services.',
         'Simplicity above all. Large text, minimal steps, voice-friendly.'),
        ('👩‍💼 Social Worker',
         'Supports multiple clients at once. Time is extremely limited. Regularly applies on behalf of others. Knows the system, but the system still wastes their time.',
         'Bulk apply, case management, proxy mode, status tracking.'),
        ('👨 Low-Income Adult',
         'No car. Lives primarily on a smartphone. Limited data plan. Limited time and energy for complex processes. May be ESL.',
         'Mobile-first, fast, no re-entry, multilingual-friendly.'),
    ]
    for name, desc, needs in personas:
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        doc.add_paragraph(desc)
        p2 = doc.add_paragraph()
        p2.add_run('Needs: ').bold = True
        p2.add_run(needs)
        doc.add_paragraph()

    add_divider(doc)

    # Journey
    doc.add_heading('User Journey: Now vs. What You\'ll Design', 1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❌ Today').bold = True
    today = [
        'Search online → Find a Ride not indexed on Google',
        'Can\'t tell which services apply to me',
        'Call each provider → long wait times',
        'Fill multiple separate forms → same data again and again',
        'Discover missing documents mid-application → abandon',
    ]
    add_numbered(doc, today)
    add_body(doc, '→ Most people give up. The burden is entirely on the user.')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run('✅ Future — What You Build').bold = True
    future = [
        'Simple questions → instantly see which services I qualify for',
        'See required documents upfront → prepare before starting',
        'Enter data once → it flows to all eligible services',
        'Apply to multiple services at once → one click to submit',
        'Track application status → no mystery, no follow-up calls',
    ]
    add_numbered(doc, future)
    add_body(doc, '→ Zero dropout. The complexity is hidden. The user just gets a ride.')

    add_divider(doc)

    # Eligibility Logic
    doc.add_heading('Eligibility Logic', 1)
    doc.add_heading('What determines who qualifies for what', 2)
    add_bullet(doc, ['💵 Income', '🎂 Age', '♿ Disability', '📍 Location'])

    add_divider(doc)

    # Success
    doc.add_heading('Success', 1)
    doc.add_heading('If the user thinks "this is easy" — you\'ve succeeded.', 2)
    doc.add_paragraph()
    add_body(doc, 'Success Criteria:')
    add_bullet(doc, [
        'Easy to use for the least tech-savvy person in the room',
        'Reduces drop-off — fewer people abandon mid-process',
        'Hides complexity — eligibility logic and data reuse work behind the scenes',
        'More people reach services that were always available — they just couldn\'t get to them',
        'Proxy-friendly — social workers can apply on behalf of clients without extra friction',
    ])

    doc.save('C:/Users/misakin/DEV/02_personal/hsi-2026-hackpack/hopelink/hackpack-hopelink.docx')
    print('Hopelink done')


make_bgc()
make_hopelink()
